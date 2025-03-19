import os
import logging
import google.generativeai as genai
import PyPDF2
import docx
import io
from fpdf import FPDF

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

GOOGLE_API_KEY = "AIzaSyDFpefCyBiSR59qtD2UVnP9WIjT7fdhF3M"  # Replace with your API key
genai.configure(api_key=GOOGLE_API_KEY)
MODEL_NAME = "models/gemini-1.5-pro-latest"

def extract_text_from_pdf(pdf_path):
    try:
        with open(pdf_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            text = "\n".join([page.extract_text() for page in reader.pages if page.extract_text()])
        return text
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        return None

def extract_text_from_docx(docx_path):
    try:
        doc = docx.Document(docx_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        return None

def extract_text_from_txt(txt_path):
    try:
        with open(txt_path, "r", encoding="utf-8") as file:
            return file.read()
    except Exception as e:
        logger.error(f"TXT extraction error: {e}")
        return None

def summarize_text(text):
    try:
        model = genai.GenerativeModel(MODEL_NAME)
        response = model.generate_content(
            f"Summarize the following text clearly and concisely:\n\n{text}",
            generation_config={"max_output_tokens": 500}
        )
        return response.text if response else None
    except Exception as e:
        logger.error(f"Summarization error: {e}")
        return None

def generate_mcqs(summary, num_questions=30):
    try:
        model = genai.GenerativeModel(MODEL_NAME)
        prompt = f"""Generate {num_questions} multiple-choice questions with four options and clear correct answers based on:
        {summary}
        Format each question as:
        Question X: [text]
        A) [option]
        B) [option]
        C) [option]
        D) [option]
        Correct Answer: [Letter]"""
        
        response = model.generate_content(prompt, generation_config={"max_output_tokens": 1500})
        return response.text if response else None
    except Exception as e:
        logger.error(f"MCQ generation error: {e}")
        return None

def save_results_to_memory(summary, mcqs, output_format):
    if output_format == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, f"Summary:\n\n{summary}\n\nMCQs:\n{mcqs}")
        buffer = io.BytesIO()
        pdf.output(buffer)
        buffer.seek(0)
        return buffer
        
    elif output_format == "docx":
        doc = docx.Document()
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(summary)
        doc.add_heading('Multiple Choice Questions', level=1)
        doc.add_paragraph(mcqs)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
        
    elif output_format == "csv":
        buffer = io.StringIO()
        writer = csv.writer(buffer)
        writer.writerow(["Summary", "MCQs"])
        writer.writerow([summary, mcqs])
        buffer.seek(0)
        return io.BytesIO(buffer.getvalue().encode('utf-8'))
        
    else:
        raise ValueError("Unsupported file format")
